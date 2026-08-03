"""Geração do relatório de visita em .docx (insumo para o PGR).

O documento é o que o técnico interno leva para o sistema de SST: dados do
chamado, setores com cargos e fotos, e as assinaturas colhidas no local.
"""
import io
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.models.catalogo import epis_do_cargo, riscos_agrupados_do_cargo
from app.models.chamado import Chamado
from app.models.enums import ROTULO_TIPO_VISITA
from app.models.setor import Setor
from app.utils.formatacao import coord_br, data_br, dt_br

AZUL_MEDSEST = RGBColor(0x1A, 0x3A, 0x5C)
LARGURA_FOTO = Inches(2.6)
LARGURA_ASSINATURA = Inches(2.2)


def _medida(valor, unidade: str) -> str | None:
    """Formata uma medição (Decimal) como '87 dB(A)', sem zeros à toa. None se vazio."""
    if valor is None:
        return None
    return f"{valor:g} {unidade}"


def _arquivo(caminho_relativo: str | None) -> Path | None:
    """Resolve o caminho gravado no banco, se o arquivo existir de fato."""
    if not caminho_relativo:
        return None
    caminho = Path(settings.UPLOAD_DIR) / caminho_relativo
    return caminho if caminho.exists() else None


def _titulo(doc: Document, texto: str, tamanho: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(tamanho)
    run.font.color.rgb = AZUL_MEDSEST


def _campo(doc: Document, rotulo: str, valor: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{rotulo}: ")
    run.bold = True
    p.add_run(valor or "-")


def _subcampo(doc: Document, rotulo: str, valor: str) -> None:
    """Campo recuado, para os detalhes de um cargo."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{rotulo}: ")
    run.bold = True
    run.font.size = Pt(10)
    v = p.add_run(valor or "-")
    v.font.size = Pt(10)


def _bloco_cargo(doc: Document, cargo) -> None:
    """Renderiza um cargo com todos os campos de PGR: nº de trabalhadores,
    jornada, riscos (agrupados por categoria) e EPIs."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    nome = p.add_run(f"• {cargo.nome_cargo}")
    nome.bold = True

    if cargo.descricao_funcao:
        _subcampo(doc, "Descrição da função", cargo.descricao_funcao)
    dados = []
    if cargo.num_trabalhadores is not None:
        dados.append(f"Nº de trabalhadores: {cargo.num_trabalhadores}")
    if cargo.jornada:
        dados.append(f"Jornada: {cargo.jornada}")
    if dados:
        _subcampo(doc, "Dados", "  |  ".join(dados))

    # Riscos: a ausência declarada é uma informação (diferente de não preenchido).
    if cargo.possui_riscos is False:
        _subcampo(doc, "Riscos", "Nenhum risco identificado (declarado no local).")
    else:
        grupos = riscos_agrupados_do_cargo(cargo.riscos or [])
        if grupos:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run("Riscos identificados:")
            r.bold = True
            r.font.size = Pt(10)
            for rotulo_cat, agentes in grupos:
                _subcampo(doc, f"  {rotulo_cat}", ", ".join(agentes))
        if cargo.riscos_outros:
            _subcampo(doc, "  Outros riscos", cargo.riscos_outros)

    # EPIs
    if cargo.utiliza_epis is False:
        _subcampo(doc, "EPIs", "Não utiliza EPI (declarado no local).")
    else:
        itens = epis_do_cargo(cargo.epis or [])
        if itens:
            _subcampo(doc, "EPIs utilizados", ", ".join(itens))
        if cargo.epis_outros:
            _subcampo(doc, "  Outros EPIs", cargo.epis_outros)


async def gerar_relatorio_word(chamado_id: uuid.UUID, db: AsyncSession) -> tuple[bytes, str]:
    """Monta o .docx. Retorna (conteúdo, nome_do_arquivo)."""
    chamado = await db.scalar(
        select(Chamado)
        .where(Chamado.id == chamado_id)
        .options(
            selectinload(Chamado.cliente),
            selectinload(Chamado.unidade_medsest),
            selectinload(Chamado.tecnico_externo),
            selectinload(Chamado.tecnico_interno),
        )
    )
    setores = (
        await db.scalars(
            select(Setor)
            .where(Setor.chamado_id == chamado_id)
            .options(selectinload(Setor.cargos), selectinload(Setor.fotos))
            .order_by(Setor.ordem, Setor.created_at)
        )
    ).all()

    cliente = chamado.cliente
    doc = Document()

    # --- Cabeçalho ---
    cabecalho = doc.add_paragraph()
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cabecalho.add_run("MedSest - Relatório de Visita Técnica")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL_MEDSEST

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Chamado #{chamado.numero_chamado}  |  {ROTULO_TIPO_VISITA[chamado.tipo_visita]}"
    )
    sub_run.font.size = Pt(11)

    doc.add_paragraph()

    # --- Cliente ---
    _titulo(doc, "Dados do Cliente")
    if cliente:
        _campo(doc, "Razão social", cliente.razao_social)
        _campo(doc, "Nome fantasia", cliente.nome_fantasia)
        _campo(doc, "CNPJ", cliente.cnpj)
        _campo(doc, "Filial", cliente.filial)
        endereco = ", ".join(filter(None, [cliente.endereco, cliente.cidade, cliente.estado]))
        _campo(doc, "Endereço", endereco)
        _campo(doc, "CEP", cliente.cep)
        _campo(doc, "Contato", cliente.nome_contato)
        _campo(doc, "Telefone", cliente.celular_contato)
        _campo(doc, "E-mail", cliente.email_contato)

    doc.add_paragraph()

    # --- Visita ---
    _titulo(doc, "Dados da Visita")
    _campo(doc, "Unidade MedSest", chamado.unidade_medsest.nome if chamado.unidade_medsest else "-")
    _campo(doc, "Técnico externo", chamado.tecnico_externo.nome if chamado.tecnico_externo else "-")
    _campo(doc, "Técnico interno", chamado.tecnico_interno.nome if chamado.tecnico_interno else "-")
    _campo(doc, "Abertura do chamado", dt_br(chamado.dt_abertura))
    _campo(doc, "Data proposta", data_br(chamado.data_proposta))
    if chamado.data_visita_alterada:
        _campo(doc, "Data reagendada", data_br(chamado.data_visita_alterada))
    _campo(doc, "Início da visita", dt_br(chamado.dt_inicio_visita))
    _campo(doc, "Fim da visita", dt_br(chamado.dt_fim_visita))
    _campo(doc, "Geolocalização (início)", coord_br(chamado.geoloc_latitude, chamado.geoloc_longitude))
    if chamado.recomendacoes:
        _campo(doc, "Recomendações", chamado.recomendacoes)

    # --- Setores ---
    doc.add_page_break()
    _titulo(doc, "Setores e Ambientes", tamanho=16)

    if not setores:
        doc.add_paragraph("Nenhum setor registrado.")

    for indice, setor in enumerate(setores, start=1):
        _titulo(doc, f"{indice}. {setor.nome}", tamanho=13)
        if setor.descricao_ambiente:
            _campo(doc, "Descrição do ambiente", setor.descricao_ambiente)
        if setor.maquinas:
            _campo(doc, "Máquinas e equipamentos", setor.maquinas)

        # Medições ambientais (só as informadas)
        medicoes = [
            f"{rotulo}: {texto}"
            for rotulo, texto in (
                ("Ruído", _medida(setor.ruido_db, "dB(A)")),
                ("Calor", _medida(setor.calor_ibutg, "°C (IBUTG)")),
                ("Iluminância", _medida(setor.iluminancia_lux, "lux")),
            )
            if texto
        ]
        if medicoes:
            _campo(doc, "Medições", "  |  ".join(medicoes))

        # Um bloco por cargo: comporta riscos, EPIs e os demais campos de PGR.
        if setor.cargos:
            p = doc.add_paragraph()
            p.add_run("Cargos / Funções").bold = True
            for cargo in setor.cargos:
                _bloco_cargo(doc, cargo)
        else:
            doc.add_paragraph("Nenhum cargo registrado neste setor.")

        # Fotos
        if setor.fotos:
            p = doc.add_paragraph()
            p.add_run("Registro fotográfico").bold = True
            for foto in setor.fotos:
                arquivo = _arquivo(foto.caminho_arquivo)
                if arquivo is None:
                    # O registro existe mas o arquivo sumiu do disco: avisa em
                    # vez de quebrar a exportação inteira.
                    doc.add_paragraph(f"[foto indisponível: {foto.nome_original or foto.caminho_arquivo}]")
                    continue
                doc.add_picture(str(arquivo), width=LARGURA_FOTO)
                legenda = doc.add_paragraph(foto.descricao or foto.nome_original or "")
                legenda.paragraph_format.space_after = Pt(8)
                for r in legenda.runs:
                    r.font.size = Pt(9)
                    r.italic = True

        doc.add_paragraph()

    # --- Assinaturas ---
    doc.add_page_break()
    _titulo(doc, "Conferência e Assinaturas", tamanho=16)
    doc.add_paragraph(
        "As informações deste relatório foram conferidas no local, ao final da "
        "visita, e assinadas pelo cliente e pelo técnico responsável."
    )
    doc.add_paragraph()

    tabela = doc.add_table(rows=1, cols=2)
    col_cliente, col_tecnico = tabela.rows[0].cells

    col_cliente.paragraphs[0].add_run("CLIENTE").bold = True
    assinatura_cliente = _arquivo(chamado.assinatura_cliente_caminho)
    if assinatura_cliente:
        col_cliente.add_paragraph().add_run().add_picture(
            str(assinatura_cliente), width=LARGURA_ASSINATURA
        )
    col_cliente.add_paragraph(f"Nome: {chamado.assinatura_cliente_nome or '-'}")
    col_cliente.add_paragraph(f"CPF: {chamado.assinatura_cliente_cpf or '-'}")
    col_cliente.add_paragraph(f"Data/hora: {dt_br(chamado.dt_assinatura_cliente)}")

    col_tecnico.paragraphs[0].add_run("TÉCNICO MEDSEST").bold = True
    assinatura_tecnico = _arquivo(chamado.assinatura_tecnico_caminho)
    if assinatura_tecnico:
        col_tecnico.add_paragraph().add_run().add_picture(
            str(assinatura_tecnico), width=LARGURA_ASSINATURA
        )
    col_tecnico.add_paragraph(f"Nome: {chamado.tecnico_externo.nome if chamado.tecnico_externo else '-'}")
    col_tecnico.add_paragraph(f"Data/hora: {dt_br(chamado.dt_assinatura_tecnico)}")

    doc.add_paragraph()
    _campo(
        doc,
        "Geolocalização das assinaturas",
        coord_br(chamado.geoloc_assinatura_latitude, chamado.geoloc_assinatura_longitude),
    )

    rodape = doc.add_paragraph()
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rodape.add_run(
        "Documento gerado pelo MedSest Visita. Este relatório reúne os insumos "
        "coletados em campo e não constitui, por si só, o PGR."
    )
    r.font.size = Pt(8)
    r.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)

    nome_cliente = (cliente.razao_social if cliente else "cliente")[:40]
    seguro = "".join(ch if ch.isalnum() or ch in " -_" else "" for ch in nome_cliente).strip()
    nome_arquivo = f"visita_{chamado.numero_chamado}_{seguro.replace(' ', '_')}.docx"

    return buffer.getvalue(), nome_arquivo
